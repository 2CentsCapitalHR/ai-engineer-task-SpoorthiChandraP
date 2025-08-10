import streamlit as st
import docx
import json
import os
import io
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ======== ADGM CHECKLIST ========
REQUIRED_DOCS = {
    "Company Incorporation": [
        "Articles of Association",
        "Memorandum of Association",
        "Board Resolution",
        "UBO Declaration Form",
        "Register of Members and Directors",
        "Incorporation Application Form"
    ],
    "Employment & HR": [
        "Standard Employment Contract Template (2024 update)",
        "Standard Employment Contract Template (2019 short version)"
    ],
    "Data Protection": [
        "Appropriate Policy Document Template"
    ],
    "Compliance & Filings": [
        "Annual Accounts & Filings"
    ]
}

# ======== Detect document type ========
def detect_doc_type(doc):
    text = "\n".join([p.text for p in doc.paragraphs]).lower()
    for process, docs in REQUIRED_DOCS.items():
        for d in docs:
            if d.lower() in text:
                return d
    return "Unknown Document"

# ======== Simple red flag checks ========
def check_red_flags(doc):
    issues = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.lower()
        if "uae federal court" in text:
            issues.append({
                "section": f"Paragraph {i+1}",
                "issue": "Mentions UAE Federal Court instead of ADGM Courts",
                "severity": "High",
                "suggestion": "Change to ADGM Courts"
            })
        if "signature" not in text and i == len(doc.paragraphs) - 1:
            issues.append({
                "section": f"Paragraph {i+1}",
                "issue": "Missing signature section",
                "severity": "Medium",
                "suggestion": "Add signature block"
            })
    return issues

# ======== Add comment (low-level XML) ========
def add_comment(paragraph, comment_text):
    try:
        start = OxmlElement('w:commentRangeStart')
        start.set(qn('w:id'), '0')
        paragraph._p.addprevious(start)

        end = OxmlElement('w:commentRangeEnd')
        end.set(qn('w:id'), '0')
        paragraph._p.addnext(end)

        comment = OxmlElement('w:comment')
        comment.set(qn('w:id'), '0')
        comment.set(qn('w:author'), 'ADGM Corporate Agent')
        comment.set(qn('w:date'), '2025-08-10T00:00:00Z')

        p = OxmlElement('w:p')
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = comment_text
        r.append(t)
        p.append(r)
        comment.append(p)

        comments_part = paragraph.part.comments_part
        comments_part.element.append(comment)
    except Exception:
        # If adding comment fails, append italic text in paragraph
        paragraph.add_run(f" <-- COMMENT: {comment_text}").italic = True

# ======== Process uploaded documents ========
def process_docs(uploaded_files):
    detected_docs = []
    all_issues = []
    reviewed_files = []

    os.makedirs("outputs", exist_ok=True)

    for uploaded_file in uploaded_files:
        file_bytes = io.BytesIO(uploaded_file.read())
        doc = docx.Document(file_bytes)

        doc_type = detect_doc_type(doc)
        detected_docs.append(doc_type)

        issues = check_red_flags(doc)
        for issue in issues:
            issue["document"] = doc_type
            para_index = int(issue["section"].split(" ")[1]) - 1
            if 0 <= para_index < len(doc.paragraphs):
                comment_text = f"{issue['issue']} | Suggestion: {issue['suggestion']}"
                add_comment(doc.paragraphs[para_index], comment_text)
        all_issues.extend(issues)

        reviewed_filename = f"reviewed_{uploaded_file.name}"
        reviewed_path = os.path.join("outputs", reviewed_filename)
        doc.save(reviewed_path)
        reviewed_files.append(reviewed_path)

    process_name = "Company Incorporation"
    required_list = REQUIRED_DOCS[process_name]
    missing_docs = [d for d in required_list if d not in detected_docs]

    result = {
        "process": process_name,
        "documents_uploaded": len(detected_docs),
        "required_documents": len(required_list),
        "missing_document": missing_docs if missing_docs else None,
        "issues_found": all_issues
    }

    output_json_path = os.path.join("outputs", "output.json")
    with open(output_json_path, "w") as fp:
        json.dump(result, fp, indent=4)

    return result, reviewed_files, output_json_path

# ======== Streamlit UI ========
st.title("ADGM Corporate Agent")

uploaded_files = st.file_uploader(
    "Upload ADGM .docx documents",
    type="docx",
    accept_multiple_files=True
)

if uploaded_files:
    result, reviewed_files, output_json_path = process_docs(uploaded_files)

    st.subheader("Analysis Result")
    st.json(result)

    with open(output_json_path, "rb") as f:
        st.download_button("Download JSON Report", f, file_name="output.json")

    for reviewed_path in reviewed_files:
        with open(reviewed_path, "rb") as f:
            st.download_button(
                label=f"Download Reviewed {os.path.basename(reviewed_path)}",
                data=f,
                file_name=os.path.basename(reviewed_path)
            )
